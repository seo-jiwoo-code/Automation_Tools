from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn
from docx2pdf import convert
from docx.oxml.ns import nsdecls
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, RGBColor, Cm, Pt
import os

## Requirements 
# 1. AutoFit Feature like Word (Autoadjust col width to reduce space wastage) - done
# 2. Minimize Row Height - done
# 3. Pad the Cells (top, bottom, right, left) - done
# 7. Fixed spacing between header and table. - expected already
# 9. Table should within page padding - expected already

# 6. Header cannot span more than 3 rows - wont do

# 8. Table should not break over mutiple pages - todo - work on this
# How is JJ gonna tell me what is needed ?
# 4. Lists within cells should have right spacing - todo
# 5. Some Cells (Sample prep fee) need to span over multiple cells -todo 

## Other specs
# 1. Header 1 : Font / Size / center-align
# 2. Header 2 : Font / Size / left-align
# 3. Table header row : Font / Size / bg-color / cell-padding / justified
# 4. Table content row : Font / Size / bg-color / cell-padding / justified
# 

def add_picture(doc, img_path, width, height):
    # p = doc.add_paragraph()
    section = doc.sections[0]
    section.top_margin = Inches(0)
    # section.left_margin = Inches(0.01)
    # section.right_margin = Inches(0.01)
    header = section.header
    p = header.add_paragraph()
    p = p.insert_paragraph_before('')
    p.paragraph_format.right_indent = -Inches(1.0)
    p.style = doc.styles['Header']
    if p.runs:
        r = p.runs[0]
    else:
        r = p.add_run()
    r.add_picture(img_path, width=width, height=height)
    return p

def align_para(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph.alignment = alignment

def add_table(doc, data):
    num_rows = len(data)
    num_columns = len(data[0])
    table = doc.add_table(rows=num_rows, cols=num_columns)
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            if isinstance(cell_data, list) and len(cell_data) > 1:
                # add_list_to_cell_2(table.cell(row_idx, col_idx))
                table.cell(row_idx, col_idx).text = ""
                p = table.cell(row_idx, col_idx).paragraphs[0]._element
                p.getparent().remove(p)
                p._p = p._element = None
                for ele in cell_data:
                    para = table.cell(row_idx, col_idx).add_paragraph(str(ele).strip(), style='List Bullet')
            elif isinstance(cell_data, list) and len(cell_data) == 1:
                table.cell(row_idx, col_idx).text = str(cell_data[0]).strip()
            elif isinstance(cell_data, list) and len(cell_data) == 0:
                table.cell(row_idx, col_idx).text = ""
            else:
                table.cell(row_idx, col_idx).text = str(cell_data).strip()
    return table

def set_doc_font(doc):
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    return doc

def set_table_alignment(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

def set_table_colwidth(table):
    table.allow_autofit = True
    table.autofit = True
    # doc.tables[t_idx]._tblPr.xpath("./w:tblW")[0].attrib["{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"] = "auto"
    for row_idx, r_val in enumerate(table.rows):
        for cell_idx, c_val in enumerate(table.rows[row_idx].cells):
            table.rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.type = 'auto'
            table.rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.w = 0

def set_table_borders(table):
    table.style = 'Table Grid'

def set_table_cellfont(table, font_size, font_name = "Times New Roman"):
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(font_size)
                    font.name = font_name

def set_table_cellalignment(table):
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_table_rowheight(table, num_rows):
    for row_idx in range(num_rows):
        table.rows[row_idx].height_rule = WD_ROW_HEIGHT_RULE.AUTO

def set_table_width(table):
    tc = table._element
    tblPr = tc.tblPr
    tblW= OxmlElement('w:tblW')
    tblW.set(qn('w:type'), "dxa")
    tblW.set(qn('w:w'), "99%")
    tblPr.append(tblW)

def set_para_spacing(paragraph, before=0, after=0, line_spacing=1, line_spacing_rule=WD_LINE_SPACING.SINGLE):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.line_spacing_rule = line_spacing_rule

def set_table_cellpadding(table, padding):
    tc = table._element
    tblPr = tc.tblPr
    tblCellMar = OxmlElement('w:tblCellMar')
    kwargs = {}
    for m in ["left", "right", "top", "bottom"]:
        node = OxmlElement("w:{}".format(m))
        node.set(qn('w:w'), str(padding))
        node.set(qn('w:type'), 'dxa')
        tblCellMar.append(node)
    tblPr.append(tblCellMar)

def set_cell_color(cell, color):
    tblCellProperties = cell._tc.get_or_add_tcPr()
    clShading = OxmlElement('w:shd')
    clShading.set(qn('w:fill'), color) #Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
    tblCellProperties.append(clShading)

def set_table_headerbgcolor(table, bgcolor = "00519E", header_type="row"):
    if header_type=="row":
        for cell_idx, c_val in enumerate(table.rows[0].cells):
            tblCellProperties = table.rows[0].cells[cell_idx]._tc.get_or_add_tcPr()
            clShading = OxmlElement('w:shd')
            clShading.set(qn('w:fill'), bgcolor) #Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
            tblCellProperties.append(clShading)
    elif header_type=="col":
        for cell_idx, c_val in enumerate(table.columns[0].cells):
            tblCellProperties = table.columns[0].cells[cell_idx]._tc.get_or_add_tcPr()
            clShading = OxmlElement('w:shd')
            clShading.set(qn('w:fill'), bgcolor) #Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
            tblCellProperties.append(clShading)

def set_table_header_italic_bold(table, header_type="row"):
    if header_type=="row":
        cell_list = table.rows[0]
    elif header_type=="col":
        cell_list = table.columns[0]
    
    for cell in cell_list.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            set_text_italic(paragraph)
            set_text_bold(paragraph)

def add_heading(doc, heading_level, heading_text):
    # Define a custom style for the heading
    heading = doc.add_heading(level=1)
    heading.text = str(heading_text).strip().replace("\n", "")
    return heading

def set_text_font(paragraph, font_name, font_size):
    paragraph.runs[0].font.name = 'Times New Roman'
    paragraph.runs[0].font.size = Pt(font_size) 
    paragraph.runs[0].font.color.rgb = RGBColor(0, 0, 0) # black

def set_text_alignment(paragraph, left_indent=0, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph.paragraph_format.alignment = alignment
    paragraph.paragraph_format.left_indent = Pt(0)

def set_text_underline(paragraph):
    paragraph.runs[0].underline = WD_UNDERLINE.SINGLE

def set_text_italic(paragraph):
    paragraph.runs[0].italic = True

def set_text_bold(paragraph):
    paragraph.runs[0].bold = True

#TO DO
def prevent_table_rowbreak(table, num_rows):
    for row in table.rows:
        row._tr.trPr.insert(0, OxmlElement('w:cantSplit'))
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if paragraph._p.pPr is None:
                    paragraph._p.insert(0, OxmlElement('w:pPr'))
                node = OxmlElement("w:keepNext")
                # node.set(qn('w:val'), "1")
                paragraph._p.pPr.append(node)

def add_story(doc, heading1, heading2, data):
    heading1 = add_heading(doc, 1, "Lab Test Details")
    set_text_font(heading1, font_name = "Times New Roman", font_size = 15)
    set_text_alignment(heading1, left_indent=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_text_underline(heading1)

    # print(heading1._p.xml)

    heading2 = add_heading(doc, 3, "Minerals per Sample")
    set_text_font(heading2, font_name = "Times New Roman", font_size = 12)
    set_text_alignment(heading2, left_indent=0, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    table = add_table(doc, data)
    num_rows = len(data)

    set_table_colwidth(table)
    set_table_borders(table)
    set_table_rowheight(table, num_rows)
    set_table_cellpadding(table, padding=200)
    set_table_headerbgcolor(table, bgcolor = 'C7C7C7')
    set_table_cellfont(table, 8)
    prevent_table_rowbreak(table, num_rows)

def add_doc_borders(doc):
    sec_pr = doc.sections[0]._sectPr # get the section properties el
    # create new borders el
    pg_borders = OxmlElement('w:pgBorders')
    # specifies how the relative positioning of the borders should be calculated
    pg_borders.set(qn('w:offsetFrom'), 'page')
    for border_name in ('top', 'left', 'bottom', 'right',): # set all borders
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single') # a single line
        border_el.set(qn('w:sz'), '4') # for meaning of  remaining attrs please look docs
        border_el.set(qn('w:space'), '24')
        border_el.set(qn('w:color'), 'auto')
        pg_borders.append(border_el) # register single border to border el
    sec_pr.append(pg_borders) # apply border changes to section

def add_doc_lib(doc):
    doc_ele = doc._element
    doc_ele.set(qn('w:cx'), "AAaaaaaaaaaaaaaaa")

def create_docx(data, filename):
    doc = Document()
    # print(doc._element.xml)

    add_doc_lib(doc)

    add_doc_borders(doc)

    add_story(doc, 
        heading1="Lab Test Details",
        heading2="Minerals per Sample",
        data=data
    )

    add_story(doc, 
        heading1="Lab Test Details",
        heading2="Minerals per Sample",
        data=data
    )

    add_story(doc, 
        heading1="Lab Test Details",
        heading2="Minerals per Sample",
        data=data
    )

    add_story(doc, 
        heading1="Lab Test Details",
        heading2="Minerals per Sample",
        data=data
    )

    # print(table._tbl.xml)
    # print(table.allow_autofit)
    # set_table_headerbgcolor(table, bgcolor = "00519E")
    doc.save(filename)

if __name__ == "__main__":
    # Sample data for the table (replace this with your actual data)
    data = [
        ["No.", "Test Parameters", "Method Reference", "TAT (calendar days)", "Min. Sample Size (g)", "Cost (exc. GST)", "Units of Measurement"],
        ["1", "Sodium (Na)", "ICP-OES, In-House Method 6324s", "7", "20", "$38", "mg/100g"],
        ["2", "Phosphate (PO₄³⁻)", "ICP-OES, In-House Method 6324s", "7", "20", "$38", "mg/100g"],
        ["2", "Phosphate (PO₄³⁻)", "ICP-OES, In-House Method 6324s", "7", "20", "$38", "mg/100g"]
#         ["2", """1. Sodium (Na)
# 2. Phosphate (PO₄³⁻)""", "ICP-OES, In-House Method 6324s", "7", "20", "$38", "mg/100g"]
    ]

    # Output file name for the Word document
    output_filename = "test_table_2.docx"
    output_filename_pdf = "test_table_2.pdf"

    create_docx(data, output_filename)

    # # Font settings for the Word document
    # font_name = "Times New Roman"
    # font_size = 12

    # # Maximum width for each column (in EMU units)
    # max_column_width = [11000, 6000, 8000]

    # # Generate the Word document with the table using custom font settings and column widths
    # # create_docx_with_table(data, output_filename, font_name, font_size, max_column_width)  
    in_file = os.path.abspath(output_filename)
    out_file = os.path.abspath(output_filename.replace(".docx", ".pdf"))

    # # Use docx2pdf library to convert Word to PDF
    convert(in_file, out_file)